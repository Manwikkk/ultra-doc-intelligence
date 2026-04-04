"""
ingestor.py — Document parsing, cleaning, and intelligent chunking.

Supports: PDF (PyMuPDF primary, pdfplumber fallback), DOCX, TXT
Chunking strategy:
  - Token-aware sliding window (500–700 tokens, ~100 overlap)
  - Section-boundary-aware: preserves Pickup / Drop / Rate / Instructions blocks
"""
from __future__ import annotations

import re
import uuid
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# ── Section keywords that should start a new chunk boundary ──────────────────
SECTION_KEYWORDS = [
    r"pickup",
    r"pick[\s\-]?up",
    r"drop[\s\-]?off",
    r"delivery",
    r"consignee",
    r"shipper",
    r"rate breakdown",
    r"freight charges",
    r"special instructions",
    r"bill of lading",
    r"pro number",
    r"carrier",
    r"equipment",
    r"weight",
    r"hazmat",
    r"commodity",
]
SECTION_PATTERN = re.compile(
    r"(?i)^.*(" + "|".join(SECTION_KEYWORDS) + r").*$",
    re.MULTILINE,
)

# Approximate chars per token (for English text)
CHARS_PER_TOKEN = 4


def _chunk_size_chars(chunk_size_tokens: int) -> int:
    return chunk_size_tokens * CHARS_PER_TOKEN


def parse_pdf(file_path: Path) -> str:
    """Extract text from PDF using PyMuPDF with pdfplumber fallback."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(file_path))
        pages = []
        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                pages.append(f"[Page {page_num + 1}]\n{text}")
        doc.close()
        full_text = "\n\n".join(pages)
        if len(full_text.strip()) > 50:
            logger.info("PyMuPDF extracted %d characters", len(full_text))
            return full_text
        raise ValueError("PyMuPDF returned insufficient text")
    except Exception as e:
        logger.warning("PyMuPDF failed (%s), falling back to pdfplumber", e)
        import pdfplumber
        pages = []
        with pdfplumber.open(str(file_path)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"[Page {i + 1}]\n{text}")
        return "\n\n".join(pages)


def parse_docx(file_path: Path) -> str:
    """Extract text from DOCX, preserving table content."""
    from docx import Document
    doc = Document(str(file_path))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def parse_txt(file_path: Path) -> str:
    """Read plain text file."""
    return file_path.read_text(encoding="utf-8", errors="replace")


def clean_text(text: str) -> str:
    """Normalize whitespace, remove noise characters."""
    # Remove null bytes and non-printable chars (keep newlines/tabs)
    text = re.sub(r"[^\x09\x0a\x0d\x20-\x7e\u00a0-\ufffd]", " ", text)
    # Collapse multiple blank lines → max 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Normalize spaces within lines
    lines = []
    for line in text.splitlines():
        line = re.sub(r"[ \t]+", " ", line).strip()
        lines.append(line)
    return "\n".join(lines)


def _find_section_boundaries(text: str) -> list[int]:
    """Return character positions of section header lines."""
    boundaries = [0]
    for match in SECTION_PATTERN.finditer(text):
        pos = match.start()
        # Only add if significantly far from last boundary
        if pos - boundaries[-1] > 200:
            boundaries.append(pos)
    return boundaries


def chunk_text(
    text: str,
    chunk_size: int = 600,
    overlap: int = 100,
) -> list[dict]:
    """
    Intelligent chunking:
    1. First split on section boundaries (semantic sections)
    2. Then apply sliding-window within sections that are too large
    3. Merge sections that are too small

    Returns list of dicts: {text, chunk_index, page}
    """
    chunk_size_chars = _chunk_size_chars(chunk_size)
    overlap_chars = _chunk_size_chars(overlap)

    # Extract page tags for metadata
    page_map: dict[int, int] = {}  # char_offset → page_number
    page_pattern = re.compile(r"\[Page (\d+)\]")
    clean_parts = []
    current_pos = 0
    last_page = 1

    for match in page_pattern.finditer(text):
        before = text[current_pos : match.start()]
        clean_parts.append(before)
        last_page = int(match.group(1))
        page_map[sum(len(p) for p in clean_parts)] = last_page
        current_pos = match.end()
    clean_parts.append(text[current_pos:])
    clean_text_no_tags = "".join(clean_parts)

    def get_page(char_offset: int) -> int:
        """Find page number for a given char offset."""
        page = 1
        for offset, pg in sorted(page_map.items()):
            if offset <= char_offset:
                page = pg
            else:
                break
        return page

    # ── Section-boundary split ────────────────────────────────────────────────
    boundaries = _find_section_boundaries(clean_text_no_tags)
    boundaries.append(len(clean_text_no_tags))

    raw_sections = []
    for i in range(len(boundaries) - 1):
        section = clean_text_no_tags[boundaries[i] : boundaries[i + 1]].strip()
        if section:
            raw_sections.append((section, boundaries[i]))

    # ── Sliding window within oversized sections ──────────────────────────────
    final_chunks = []
    chunk_idx = 0

    for section_text, section_start in raw_sections:
        if len(section_text) <= chunk_size_chars:
            page = get_page(section_start)
            final_chunks.append({
                "text": section_text,
                "chunk_index": chunk_idx,
                "page": page,
            })
            chunk_idx += 1
        else:
            # Sliding window
            start = 0
            while start < len(section_text):
                end = min(start + chunk_size_chars, len(section_text))
                chunk = section_text[start:end].strip()
                if chunk:
                    page = get_page(section_start + start)
                    final_chunks.append({
                        "text": chunk,
                        "chunk_index": chunk_idx,
                        "page": page,
                    })
                    chunk_idx += 1
                if end == len(section_text):
                    break
                start += chunk_size_chars - overlap_chars

    # ── Merge tiny chunks into predecessors ──────────────────────────────────
    MIN_CHUNK_CHARS = 150
    merged = []
    for chunk in final_chunks:
        if merged and len(chunk["text"]) < MIN_CHUNK_CHARS:
            merged[-1]["text"] += " " + chunk["text"]
        else:
            merged.append(chunk)

    # Re-index after merge
    for i, c in enumerate(merged):
        c["chunk_index"] = i

    logger.info("Chunked into %d chunks", len(merged))
    return merged


def ingest_file(file_path: Path, chunk_size: int = 600, overlap: int = 100) -> dict:
    """
    Full ingestion pipeline: parse → clean → chunk.

    Returns:
        {
            doc_id: str,
            filename: str,
            chunks: list[{text, chunk_index, page}],
            raw_text: str,
        }
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        raw_text = parse_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        raw_text = parse_docx(file_path)
    elif suffix == ".txt":
        raw_text = parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    cleaned = clean_text(raw_text)
    chunks = chunk_text(cleaned, chunk_size=chunk_size, overlap=overlap)

    doc_id = str(uuid.uuid4())
    logger.info(
        "Ingested '%s' → doc_id=%s, %d chunks",
        file_path.name,
        doc_id,
        len(chunks),
    )

    return {
        "doc_id": doc_id,
        "filename": file_path.name,
        "chunks": chunks,
        "raw_text": cleaned,
    }
