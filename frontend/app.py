"""
app.py — Ultra Doc-Intelligence v3 Frontend
Redesigned: Amber/slate luxury terminal aesthetic.
New feature: Shipment Event Timeline — auto-generated from extraction data.
"""
from __future__ import annotations

import io
import os
import re
import uuid
import base64
import logging
import datetime
from pathlib import Path

import numpy as np
import requests
import streamlit as st
import streamlit.components.v1 as components

logger = logging.getLogger(__name__)


def get_favicon():
    icon_path = Path(__file__).parent / "Document.png"
    if icon_path.exists():
        with open(icon_path, "rb") as f:
            data = base64.b64encode(f.read()).decode()
        return f"data:image/png;base64,{data}"
    return "📄"


st.set_page_config(
    page_title="Ultra Doc-Intelligence",
    page_icon=get_favicon(),
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ══════════════════════════════════════════════════════════════════════════════
#  CSS — Amber/Slate luxury terminal
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Mono:ital,wght@0,300;0,400;0,500;1,400&family=Instrument+Serif:ital@0;1&family=DM+Sans:wght@300;400;500;600&display=swap');

:root {
  --bg:         #0d1117;
  --bg-2:       #131920;
  --bg-3:       #1a2230;
  --border:     rgba(255,255,255,0.07);
  --border-2:   rgba(255,255,255,0.12);
  --amber:      #f6a623;
  --amber-dim:  #c47e0d;
  --amber-glow: rgba(246,166,35,0.15);
  --amber-bg:   rgba(246,166,35,0.06);
  --text:       #e8dcc8;
  --text-dim:   rgba(232,220,200,0.45);
  --text-faint: rgba(232,220,200,0.2);
  --green:      #4ade80;
  --red:        #f87171;
  --blue:       #60a5fa;
  --mono:       'DM Mono', 'Courier New', monospace;
  --serif:      'Instrument Serif', Georgia, serif;
  --sans:       'DM Sans', system-ui, sans-serif;
}

html, body, [class*="css"] {
  font-family: var(--sans) !important;
  color: var(--text) !important;
}
.stApp { background: var(--bg) !important; }
.stApp > header { background: transparent !important; }
#MainMenu, footer { visibility: hidden; }

.block-container {
  padding-top: 0 !important;
  padding-bottom: 2rem !important;
  max-width: 100% !important;
  padding-left: 0 !important;
  padding-right: 0 !important;
}
section[data-testid="stSidebar"] { display: none !important; }

/* ── Top bar ── */
.topbar {
  background: var(--bg-2);
  border-bottom: 1px solid var(--border);
  padding: 0 2rem;
  display: flex;
  align-items: center;
  justify-content: space-between;
  height: 56px;
}
.topbar-left { display: flex; align-items: center; gap: 14px; }
.topbar-logo {
  font-family: var(--serif);
  font-size: 1.15rem;
  color: var(--amber);
  letter-spacing: 0.02em;
  font-style: italic;
}
.topbar-div { width:1px; height:18px; background:var(--border-2); }
.topbar-tag {
  font-family: var(--mono);
  font-size: 0.6rem;
  color: var(--text-faint);
  letter-spacing: 1.5px;
  text-transform: uppercase;
}
.status-pill {
  display: inline-flex; align-items: center; gap: 7px;
  font-family: var(--mono);
  font-size: 0.65rem;
  color: var(--text-dim);
  letter-spacing: 0.3px;
  background: var(--bg-3);
  border: 1px solid var(--border);
  border-radius: 6px;
  padding: 4px 10px;
}
.dot { width:6px; height:6px; border-radius:50%; flex-shrink:0; }
.dot-on  { background: var(--green); box-shadow: 0 0 6px rgba(74,222,128,0.7); }
.dot-off { background: var(--red);   box-shadow: 0 0 5px rgba(248,113,113,0.5); }

/* ── Panel labels ── */
.panel-label {
  font-family: var(--mono);
  font-size: 0.57rem;
  font-weight: 500;
  letter-spacing: 2px;
  text-transform: uppercase;
  color: var(--amber-dim);
  margin-bottom: 0.6rem;
  display: flex; align-items: center; gap: 8px;
}
.panel-label::after {
  content: ''; flex: 1; height: 1px; background: var(--border);
}
.panel-divider {
  border: none; border-top: 1px solid var(--border); margin: 1.1rem 0;
}

/* ── Doc chip ── */
.doc-chip {
  background: var(--amber-bg);
  border: 1px solid rgba(246,166,35,0.2);
  border-radius: 8px; padding: 9px 12px; margin-top: 0.4rem;
}
.doc-chip-name {
  font-size: 0.8rem; color: var(--amber); font-weight: 500;
  white-space: nowrap; overflow: hidden; text-overflow: ellipsis;
}
.doc-chip-id {
  font-family: var(--mono); font-size: 0.58rem;
  color: var(--text-faint); margin-top: 3px;
}

/* ── How-to ── */
.how-card {
  background: var(--bg-3); border: 1px solid var(--border);
  border-radius: 9px; padding: 0.9rem 1rem; margin-top: 0.6rem;
}
.how-step {
  display: flex; align-items: flex-start; gap: 9px;
  margin-bottom: 0.55rem; font-size: 0.77rem;
  color: var(--text-dim); line-height: 1.5;
}
.how-step:last-child { margin-bottom: 0; }
.how-num {
  font-family: var(--mono); font-size: 0.6rem; color: var(--amber-dim);
  background: rgba(246,166,35,0.08); border: 1px solid rgba(246,166,35,0.15);
  border-radius: 4px; padding: 1px 6px; flex-shrink: 0; margin-top: 1px;
}

/* ── Offline banner ── */
.offline-banner {
  background: rgba(248,113,113,0.05);
  border: 1px solid rgba(248,113,113,0.15);
  border-left: 3px solid var(--red);
  border-radius: 8px; padding: 0.8rem 1rem; margin-bottom: 1.1rem;
}
.offline-title {
  font-family: var(--mono); font-size: 0.72rem; color: var(--red);
  margin-bottom: 5px; letter-spacing: 0.5px;
}
.offline-body { font-size: 0.75rem; color: rgba(248,113,113,0.6); line-height: 1.7; }
.offline-cmd {
  font-family: var(--mono); font-size: 0.68rem;
  background: rgba(0,0,0,0.3); color: #fca5a5;
  border-radius: 4px; padding: 3px 8px; display: inline-block; margin-top: 4px;
}

/* ── Chat ── */
.chat-wrap { display: flex; flex-direction: column; gap: 1.1rem; }
.chat-q-row { display: flex; justify-content: flex-end; margin-bottom: 0.45rem; }
.chat-q {
  background: linear-gradient(135deg, rgba(246,166,35,0.16), rgba(246,166,35,0.07));
  border: 1px solid rgba(246,166,35,0.22);
  color: var(--amber); border-radius: 13px 13px 3px 13px;
  padding: 0.6rem 0.95rem; max-width: 88%;
  font-size: 0.875rem; line-height: 1.5; font-weight: 500;
}
.chat-a-row { display: flex; justify-content: flex-start; margin-bottom: 0.35rem; }
.chat-a {
  background: var(--bg-2); border: 1px solid var(--border);
  border-radius: 3px 13px 13px 13px;
  padding: 0.8rem 1rem; max-width: 92%;
  font-size: 0.87rem; color: var(--text); line-height: 1.8;
}
.chat-a-guard {
  background: rgba(248,113,113,0.05); border: 1px solid rgba(248,113,113,0.15);
  border-left: 3px solid var(--red); border-radius: 3px 13px 13px 13px;
  padding: 0.8rem 1rem; max-width: 92%;
  font-size: 0.87rem; color: #fca5a5; line-height: 1.7;
}
.chat-meta {
  display: flex; align-items: center; gap: 7px;
  margin-top: 0.4rem; flex-wrap: wrap;
}
.chat-ts { font-family: var(--mono); font-size: 0.6rem; color: var(--text-faint); }

/* ── Badges ── */
.badge {
  display: inline-flex; align-items: center; gap: 3px;
  padding: 2px 7px; border-radius: 999px;
  font-family: var(--mono); font-size: 0.58rem; font-weight: 500; letter-spacing: 0.3px;
}
.badge-high   { background: rgba(74,222,128,0.1);  color: var(--green); border: 1px solid rgba(74,222,128,0.2); }
.badge-med    { background: rgba(246,166,35,0.1);  color: var(--amber); border: 1px solid rgba(246,166,35,0.2); }
.badge-low    { background: rgba(248,113,113,0.1); color: var(--red);   border: 1px solid rgba(248,113,113,0.2); }
.badge-groq   { background: rgba(74,222,128,0.08); color: #4ade80; border: 1px solid rgba(74,222,128,0.15); }
.badge-openai { background: rgba(96,165,250,0.08); color: var(--blue);  border: 1px solid rgba(96,165,250,0.15); }
.badge-gemini { background: rgba(96,165,250,0.08); color: var(--blue);  border: 1px solid rgba(96,165,250,0.15); }
.badge-ollama { background: rgba(246,166,35,0.08); color: var(--amber); border: 1px solid rgba(246,166,35,0.15); }
.badge-none   { background: rgba(248,113,113,0.08);color: var(--red);   border: 1px solid rgba(248,113,113,0.15); }

.turn-divider {
  border: none; border-top: 1px solid var(--border); margin: 0.2rem 0 0 0; opacity: 0.4;
}

/* ── Source chunk ── */
.src-chunk {
  background: var(--bg-3); border: 1px solid var(--border);
  border-radius: 7px; padding: 0.65rem 0.85rem; margin: 0.3rem 0;
  font-size: 0.78rem; color: var(--text-dim); line-height: 1.7;
}
.src-meta {
  font-family: var(--mono); font-size: 0.56rem; color: var(--text-faint);
  margin-bottom: 0.3rem; letter-spacing: 0.5px; text-transform: uppercase;
}
.log-ln {
  font-family: var(--mono); font-size: 0.63rem;
  color: var(--text-faint); padding: 1px 0; line-height: 1.6;
}

/* ── Chat empty ── */
.chat-empty {
  display: flex; flex-direction: column; align-items: center;
  justify-content: center; padding: 5rem 2rem; text-align: center;
}
.empty-glyph {
  font-family: var(--serif); font-size: 3.5rem; font-style: italic;
  color: rgba(246,166,35,0.1); margin-bottom: 0.85rem; line-height: 1;
}
.empty-text { font-size: 0.85rem; line-height: 1.8; color: var(--text-faint); max-width: 280px; }

/* ── Tabs ── */
.stTabs [data-baseweb="tab-list"] {
  background: var(--bg-2) !important;
  border: 1px solid var(--border) !important;
  border-radius: 9px !important; padding: 3px !important;
  gap: 2px !important; margin-bottom: 1.1rem !important;
}
.stTabs [data-baseweb="tab"] {
  background: transparent !important; border-radius: 6px !important;
  color: var(--text-dim) !important; font-family: var(--mono) !important;
  font-size: 0.68rem !important; font-weight: 500 !important;
  letter-spacing: 0.8px !important; padding: 6px 14px !important;
  border: none !important; text-transform: uppercase !important;
}
.stTabs [aria-selected="true"] {
  background: rgba(246,166,35,0.1) !important;
  color: var(--amber) !important;
  border: 1px solid rgba(246,166,35,0.2) !important;
}
.stTabs [data-baseweb="tab-panel"] { padding-top: 0.2rem !important; }

/* ── Extraction ── */
.ext-stats {
  display: grid; grid-template-columns: repeat(3,1fr);
  gap: 0.7rem; margin-bottom: 1.1rem;
}
.ext-stat {
  background: var(--bg-2); border: 1px solid var(--border);
  border-radius: 9px; padding: 0.85rem 1rem; text-align: center;
}
.ext-stat-val {
  font-family: var(--mono); font-size: 1.25rem;
  font-weight: 500; letter-spacing: -0.5px; margin-bottom: 4px;
}
.ext-stat-key {
  font-family: var(--mono); font-size: 0.55rem;
  color: var(--text-faint); text-transform: uppercase;
  letter-spacing: 1.2px;
}
.field-card {
  background: var(--bg-2); border: 1px solid var(--border);
  border-radius: 9px; overflow: hidden; margin-bottom: 0.85rem;
}
.field-row {
  display: flex; justify-content: space-between; align-items: flex-start;
  padding: 9px 13px; border-bottom: 1px solid var(--border); gap: 1.5rem;
}
.field-row:last-child { border-bottom: none; }
.field-key {
  font-family: var(--mono); font-size: 0.67rem; color: var(--text-faint);
  flex-shrink: 0; min-width: 145px; letter-spacing: 0.2px; padding-top: 1px;
}
.field-val { font-size: 0.82rem; color: var(--text); font-weight: 500; text-align: right; word-break: break-word; }
.field-null { font-family: var(--mono); font-size: 0.67rem; color: var(--text-faint); text-align: right; font-style: italic; opacity: 0.45; }
.field-val-amber { font-family: var(--mono); font-size: 0.82rem; color: var(--amber); font-weight: 500; text-align: right; }

/* ── Timeline ── */
.timeline-card {
  background: var(--bg-2); border: 1px solid var(--border);
  border-radius: 9px; padding: 1.1rem 1.25rem; margin-bottom: 0.85rem;
}
.timeline-header {
  display: flex; align-items: center; gap: 10px; margin-bottom: 1.1rem;
}
.timeline-title { font-family: var(--serif); font-size: 1rem; font-style: italic; color: var(--text); }
.timeline-badge {
  font-family: var(--mono); font-size: 0.55rem; color: var(--amber-dim);
  background: rgba(246,166,35,0.07); border: 1px solid rgba(246,166,35,0.14);
  padding: 2px 7px; border-radius: 4px; letter-spacing: 1px; text-transform: uppercase;
}
.timeline-wrap {
  position: relative; padding: 0.25rem 0 0.25rem 1.75rem;
}
.timeline-line {
  position: absolute; left: 7px; top: 0; bottom: 0; width: 1px;
  background: linear-gradient(to bottom, transparent, rgba(196,126,13,0.4), rgba(196,126,13,0.4), transparent);
}
.timeline-event { position: relative; margin-bottom: 1.3rem; }
.timeline-event:last-child { margin-bottom: 0; }
.tl-dot {
  position: absolute; left: -1.6rem; top: 4px;
  width: 9px; height: 9px; border-radius: 50%;
  border: 2px solid var(--amber-dim); background: var(--bg);
  box-shadow: 0 0 6px rgba(246,166,35,0.25);
}
.tl-dot-filled {
  position: absolute; left: -1.6rem; top: 4px;
  width: 9px; height: 9px; border-radius: 50%;
  background: var(--amber); box-shadow: 0 0 8px rgba(246,166,35,0.45);
}
.tl-label {
  font-family: var(--mono); font-size: 0.58rem; color: var(--amber-dim);
  text-transform: uppercase; letter-spacing: 1.2px; margin-bottom: 3px;
}
.tl-value { font-size: 0.875rem; color: var(--text); font-weight: 500; line-height: 1.4; }
.tl-value-dim { font-size: 0.875rem; color: var(--text-faint); font-style: italic; line-height: 1.4; }
.tl-sub { font-family: var(--mono); font-size: 0.62rem; color: var(--text-faint); margin-top: 2px; }
.tl-empty {
  font-family: var(--mono); font-size: 0.72rem; color: var(--text-faint);
  font-style: italic; text-align: center; padding: 1.5rem; opacity: 0.5;
}

/* ── Ext empty ── */
.ext-empty {
  text-align: center; padding: 4rem 2rem; color: var(--text-faint);
}
.ext-empty-glyph {
  font-family: var(--serif); font-size: 2.8rem; font-style: italic;
  color: rgba(246,166,35,0.1); margin-bottom: 0.7rem;
}
.ext-empty-text { font-size: 0.83rem; line-height: 1.8; }

/* ── Streamlit overrides ── */
[data-testid="stForm"] {
  border: none !important; padding: 0 !important;
  background: transparent !important; box-shadow: none !important;
}
[data-testid="InputInstructions"] { display: none !important; }

.stTextArea textarea {
  background: var(--bg-3) !important;
  border: 1px solid var(--border-2) !important;
  border-radius: 8px !important; color: var(--text) !important;
  font-family: var(--sans) !important; font-size: 0.875rem !important;
  padding: 0.7rem 0.9rem !important; line-height: 1.55 !important;
  caret-color: var(--amber) !important;
}
.stTextArea textarea:focus {
  border-color: rgba(246,166,35,0.35) !important;
  box-shadow: 0 0 0 3px rgba(246,166,35,0.05) !important;
}
.stTextArea textarea::placeholder { color: var(--text-faint) !important; }

.stTextInput input {
  background: var(--bg-3) !important;
  border: 1px solid var(--border-2) !important;
  border-radius: 7px !important; color: var(--text) !important;
  font-family: var(--mono) !important; font-size: 0.8rem !important;
}
.stTextInput input:focus { border-color: rgba(246,166,35,0.35) !important; }

.stButton > button, .stFormSubmitButton > button {
  font-family: var(--sans) !important; font-weight: 600 !important;
  font-size: 0.84rem !important; border-radius: 7px !important;
  transition: all 0.15s ease !important; letter-spacing: 0.1px !important;
}
.stButton > button[kind="primary"],
.stFormSubmitButton > button[kind="primaryFormSubmit"] {
  background: linear-gradient(135deg, #b87212, #f6a623) !important;
  border: none !important; color: #0d1117 !important; font-weight: 700 !important;
}
.stButton > button[kind="primary"]:hover,
.stFormSubmitButton > button[kind="primaryFormSubmit"]:hover {
  opacity: 0.88 !important; transform: translateY(-1px) !important;
  box-shadow: 0 4px 16px rgba(246,166,35,0.2) !important;
}
.stButton > button[kind="secondary"] {
  background: var(--bg-3) !important; border: 1px solid var(--border-2) !important;
  color: var(--text-dim) !important;
}
.stButton > button[kind="secondary"]:hover {
  background: rgba(246,166,35,0.05) !important;
  border-color: rgba(246,166,35,0.2) !important; color: var(--amber) !important;
}

.stProgress > div > div {
  background: linear-gradient(90deg, var(--amber-dim), var(--amber)) !important;
  border-radius: 4px !important;
}
.stProgress { background: rgba(255,255,255,0.05) !important; border-radius: 4px; }

div[data-testid="stFileUploader"] {
  background: rgba(246,166,35,0.02) !important;
  border: 1.5px dashed rgba(246,166,35,0.18) !important; border-radius: 9px !important;
}
.streamlit-expanderHeader {
  background: var(--bg-2) !important; border: 1px solid var(--border) !important;
  border-radius: 6px !important; font-family: var(--mono) !important;
  font-size: 0.65rem !important; color: var(--text-dim) !important;
  letter-spacing: 0.5px !important; text-transform: uppercase !important;
}
.stSuccess { background: rgba(74,222,128,0.06) !important; border: 1px solid rgba(74,222,128,0.15) !important; border-radius: 7px !important; color: #86efac !important; }
.stError   { background: rgba(248,113,113,0.06) !important; border: 1px solid rgba(248,113,113,0.15) !important; border-radius: 7px !important; color: #fca5a5 !important; }
.stInfo    { background: rgba(246,166,35,0.06) !important; border: 1px solid rgba(246,166,35,0.15) !important; border-radius: 7px !important; color: var(--amber) !important; }

[data-testid="column"] { padding: 0 0.5rem !important; }
[data-testid="column"]:first-child { padding-left: 0 !important; }
[data-testid="column"]:last-child  { padding-right: 0 !important; }

::-webkit-scrollbar { width: 4px; height: 4px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: rgba(246,166,35,0.12); border-radius: 2px; }
::-webkit-scrollbar-thumb:hover { background: rgba(246,166,35,0.28); }
</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
#  Embedding pipeline (unchanged from original)
# ══════════════════════════════════════════════════════════════════════════════

@st.cache_resource(show_spinner="Loading embedding model…")
def load_embedding_model():
    from sentence_transformers import SentenceTransformer
    return SentenceTransformer("BAAI/bge-small-en")


SECTION_KEYWORDS = [
    r"pickup", r"pick[\s\-]?up", r"drop[\s\-]?off", r"delivery",
    r"consignee", r"shipper", r"rate breakdown", r"freight charges",
    r"special instructions", r"bill of lading", r"pro number",
    r"carrier", r"equipment", r"weight", r"hazmat", r"commodity",
]
SECTION_PATTERN = re.compile(
    r"(?i)^.*(" + "|".join(SECTION_KEYWORDS) + r").*$", re.MULTILINE
)
CHARS_PER_TOKEN = 4
MAX_CHUNKS = 100


def _parse_pdf(file_bytes: bytes) -> str:
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for i, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                pages.append(f"[Page {i + 1}]\n{text}")
        doc.close()
        full = "\n\n".join(pages)
        if len(full.strip()) > 50:
            return full
        raise ValueError("Insufficient text")
    except Exception:
        import pdfplumber
        pages = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"[Page {i + 1}]\n{text}")
        return "\n\n".join(pages)


def _parse_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _clean_text(text: str) -> str:
    text = re.sub(r"[^\x09\x0a\x0d\x20-\x7e\u00a0-\ufffd]", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(lines)


def _chunk_text(text: str, chunk_size: int = 600, overlap: int = 100) -> list[dict]:
    chunk_size_chars = chunk_size * CHARS_PER_TOKEN
    overlap_chars    = overlap * CHARS_PER_TOKEN
    page_map: dict[int, int] = {}
    page_pattern = re.compile(r"\[Page (\d+)\]")
    clean_parts, current_pos = [], 0
    for match in page_pattern.finditer(text):
        clean_parts.append(text[current_pos:match.start()])
        page_map[sum(len(p) for p in clean_parts)] = int(match.group(1))
        current_pos = match.end()
    clean_parts.append(text[current_pos:])
    clean_text_no_tags = "".join(clean_parts)

    def get_page(offset: int) -> int:
        page = 1
        for o, p in sorted(page_map.items()):
            if o <= offset:
                page = p
            else:
                break
        return page

    boundaries = [0]
    for m in SECTION_PATTERN.finditer(clean_text_no_tags):
        if m.start() - boundaries[-1] > 200:
            boundaries.append(m.start())
    boundaries.append(len(clean_text_no_tags))

    raw_sections = [
        (clean_text_no_tags[boundaries[i]:boundaries[i+1]].strip(), boundaries[i])
        for i in range(len(boundaries)-1)
        if clean_text_no_tags[boundaries[i]:boundaries[i+1]].strip()
    ]

    final_chunks, chunk_idx = [], 0
    for section_text, section_start in raw_sections:
        if len(section_text) <= chunk_size_chars:
            final_chunks.append({"text": section_text, "chunk_index": chunk_idx, "page": get_page(section_start)})
            chunk_idx += 1
        else:
            start = 0
            while start < len(section_text):
                end   = min(start + chunk_size_chars, len(section_text))
                chunk = section_text[start:end].strip()
                if chunk:
                    final_chunks.append({"text": chunk, "chunk_index": chunk_idx, "page": get_page(section_start + start)})
                    chunk_idx += 1
                if end == len(section_text):
                    break
                start += chunk_size_chars - overlap_chars

    MIN_CHUNK_CHARS = 150
    merged = []
    for chunk in final_chunks:
        if merged and len(chunk["text"]) < MIN_CHUNK_CHARS:
            merged[-1]["text"] += " " + chunk["text"]
        else:
            merged.append(chunk)
    for i, c in enumerate(merged):
        c["chunk_index"] = i
    return merged


def process_document(file_bytes: bytes, filename: str):
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        raw_text = _parse_pdf(file_bytes)
    elif suffix in (".docx", ".doc"):
        raw_text = _parse_docx(file_bytes)
    elif suffix == ".txt":
        raw_text = file_bytes.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    cleaned    = _clean_text(raw_text)
    chunks     = _chunk_text(cleaned)
    if len(chunks) > MAX_CHUNKS:
        chunks = chunks[:MAX_CHUNKS]
        for i, c in enumerate(chunks):
            c["chunk_index"] = i

    model    = load_embedding_model()
    prefixed = [f"passage: {c['text']}" for c in chunks]
    embeddings = model.encode(
        prefixed, batch_size=32, show_progress_bar=False,
        convert_to_numpy=True, normalize_embeddings=True,
    ).astype(np.float32)
    return str(uuid.uuid4()), chunks, embeddings


def embed_query(question: str) -> np.ndarray:
    model = load_embedding_model()
    return model.encode(
        [f"query: {question}"],
        show_progress_bar=False, convert_to_numpy=True, normalize_embeddings=True,
    ).astype(np.float32)


# ══════════════════════════════════════════════════════════════════════════════
#  Session state
# ══════════════════════════════════════════════════════════════════════════════
_DEFAULTS = {
    "doc_id":         None,
    "filename":       None,
    "chat_history":   [],
    "extract_result": None,
    "active_tab":     0,
    "backend_url":    os.getenv("BACKEND_URL", "http://localhost:8000"),
    "show_url_cfg":   False,
}
for k, v in _DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ══════════════════════════════════════════════════════════════════════════════
#  Helpers
# ══════════════════════════════════════════════════════════════════════════════
def get_url() -> str:
    return st.session_state.backend_url.rstrip("/")


def backend_check() -> tuple[bool, str]:
    url = get_url()
    try:
        r = requests.get(f"{url}/health", timeout=5)
        if r.status_code == 200:
            return True, f"Connected · {url}"
        return False, f"HTTP {r.status_code} at {url}"
    except requests.exceptions.ConnectionError:
        return False, f"Connection refused at {url}"
    except requests.exceptions.Timeout:
        return False, f"Timed out at {url}"
    except Exception as e:
        return False, str(e)


def provider_html(provider: str) -> str:
    p    = provider.lower()
    css  = {"groq":"badge-groq","openai":"badge-openai","gemini":"badge-gemini","ollama":"badge-ollama"}.get(p,"badge-none")
    icon = {"groq":"⚡","openai":"◆","gemini":"✦","ollama":"🦙"}.get(p,"•")
    return f'<span class="badge {css}">{icon} {provider}</span>'


def conf_html(score: float) -> str:
    if score >= 0.70:   css, lbl = "badge-high",  f"{score:.0%} HIGH"
    elif score >= 0.45: css, lbl = "badge-med",   f"{score:.0%} MED"
    else:               css, lbl = "badge-low",   f"{score:.0%} LOW"
    return f'<span class="badge {css}">{lbl}</span>'


def fmt_time(iso: str) -> str:
    try:
        return datetime.datetime.fromisoformat(iso).strftime("%H:%M")
    except Exception:
        return ""


# ══════════════════════════════════════════════════════════════════════════════
#  NEW FEATURE: Shipment Event Timeline
# ══════════════════════════════════════════════════════════════════════════════
def render_timeline(data: dict) -> str:
    """
    Build an HTML event timeline from extracted shipment fields.
    Shows the complete load journey: Created → Origin → Pickup → Transit → Delivery.
    Filled dots = data found; hollow dots = data missing.
    """
    sid = data.get("shipment_id")
    shipper  = data.get("shipper")
    consignee = data.get("consignee")
    carrier  = data.get("carrier_name")
    pu_dt    = data.get("pickup_datetime")
    dl_dt    = data.get("delivery_datetime")
    rate     = data.get("rate")
    currency = data.get("currency", "")
    weight   = data.get("weight")
    equip    = data.get("equipment_type")
    mode     = data.get("mode")

    rate_str = f"${rate:,.2f} {currency}".strip() if rate else "—"
    meta_str = "  ·  ".join(filter(None, [
        f"Equip: {equip}" if equip else None,
        f"Mode: {mode}" if mode else None,
        f"Weight: {weight}" if weight else None,
    ])) or "—"

    def event(label: str, value: str, sub: str, filled: bool) -> str:
        dot_cls = "tl-dot-filled" if filled else "tl-dot"
        val_cls = "tl-value" if filled else "tl-value-dim"
        return f"""<div class="timeline-event">
  <div class="{dot_cls}"></div>
  <div class="tl-label">{label}</div>
  <div class="{val_cls}">{value}</div>
  <div class="tl-sub">{sub}</div>
</div>"""

    events = []
    events.append(event(
        "Load Created",
        f"Reference: {sid}" if sid else "Load Created",
        f"ID: {sid or '—'}",
        bool(sid),
    ))
    events.append(event(
        "Origin / Shipper",
        shipper if shipper else "Not extracted",
        meta_str,
        bool(shipper),
    ))
    events.append(event(
        "Pickup",
        pu_dt if pu_dt else "Date not extracted",
        f"Weight: {weight or '—'}",
        bool(pu_dt),
    ))
    events.append(event(
        "In Transit",
        carrier if carrier else "Carrier not specified",
        f"Rate: {rate_str}",
        bool(carrier),
    ))
    events.append(event(
        "Delivery",
        dl_dt if dl_dt else "Date not extracted",
        f"Consignee: {consignee or '—'}",
        bool(dl_dt),
    ))

    inner = "\n".join(events)
    return f'<div class="timeline-wrap"><div class="timeline-line"></div>{inner}</div>'


# ══════════════════════════════════════════════════════════════════════════════
#  RENDER
# ══════════════════════════════════════════════════════════════════════════════
alive, alive_detail = backend_check()

# ── Top bar ───────────────────────────────────────────────────────────────────
dot_cls = "dot-on" if alive else "dot-off"
st.markdown(f"""
<div class="topbar">
  <div class="topbar-left">
    <div class="topbar-logo">Ultra Doc-Intelligence</div>
    <div class="topbar-div"></div>
    <div class="topbar-tag">TMS · Document Analysis</div>
  </div>
  <div class="status-pill">
    <span class="dot {dot_cls}"></span>
    {"API Online" if alive else "API Offline"}
  </div>
</div>
""", unsafe_allow_html=True)

# ── URL configurator ──────────────────────────────────────────────────────────
_t1, _t2 = st.columns([1, 9])
with _t1:
    if st.button("⚙", key="toggle_url", help="Backend URL", type="secondary"):
        st.session_state.show_url_cfg = not st.session_state.show_url_cfg

if st.session_state.show_url_cfg:
    with _t2:
        _u1, _u2 = st.columns([5, 1])
        with _u1:
            new_url = st.text_input("URL", value=st.session_state.backend_url,
                                    label_visibility="collapsed", key="url_input",
                                    placeholder="http://localhost:8000")
        with _u2:
            if st.button("Apply", type="primary"):
                st.session_state.backend_url = new_url.strip().rstrip("/")
                st.rerun()
    st.markdown(
        f'<div style="font-family:var(--mono,monospace);font-size:0.62rem;'
        f'color:rgba(232,220,200,0.28);padding:0 0 0.4rem 0">{alive_detail}</div>',
        unsafe_allow_html=True,
    )

# ── Two columns ───────────────────────────────────────────────────────────────
left_col, right_col = st.columns([1, 1.8], gap="large")


# ════════════════════════════════════
#  LEFT
# ════════════════════════════════════
with left_col:

    if not alive:
        st.markdown(f"""
<div class="offline-banner">
  <div class="offline-title">// BACKEND OFFLINE</div>
  <div class="offline-body">
    {alive_detail}<br>
    <span class="offline-cmd">cd backend &amp;&amp; uvicorn main:app --host 127.0.0.1 --port 8000 --reload</span>
  </div>
</div>
""", unsafe_allow_html=True)

    st.markdown('<div class="panel-label">Document</div>', unsafe_allow_html=True)

    uploaded_file = st.file_uploader(
        "Upload",
        type=["pdf", "docx", "txt"],
        help="Bill of Lading, Rate Confirmation, Carrier RC, or any logistics document",
        label_visibility="collapsed",
        disabled=not alive,
    )

    if uploaded_file and st.session_state.filename != uploaded_file.name:
        with st.spinner(f"Parsing & embedding {uploaded_file.name}…"):
            try:
                fb = uploaded_file.getvalue()
                doc_id, chunks, embeddings = process_document(fb, uploaded_file.name)
                payload = {
                    "doc_id": doc_id, "filename": uploaded_file.name,
                    "chunks": chunks, "embeddings": embeddings.tolist(),
                }
                resp = requests.post(f"{get_url()}/upload_embeddings", json=payload, timeout=120)
                if resp.status_code == 200:
                    d = resp.json()
                    st.session_state.doc_id         = d["doc_id"]
                    st.session_state.filename       = uploaded_file.name
                    st.session_state.chat_history   = []
                    st.session_state.extract_result = None
                    st.session_state.active_tab     = 0
                    st.success(f"✓ {uploaded_file.name}  ·  {d['chunks_count']} chunks")
                else:
                    st.error(f"Upload failed: {resp.json().get('detail', resp.text)}")
            except requests.exceptions.ConnectionError:
                st.error(f"Cannot reach {get_url()}. Check ⚙ URL.")
            except Exception as e:
                st.error(f"Error: {e}")

    if st.session_state.doc_id:
        st.markdown(
            f'<div class="doc-chip">'
            f'<div class="doc-chip-name">📄 {st.session_state.filename}</div>'
            f'<div class="doc-chip-id">{st.session_state.doc_id}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    st.markdown('<hr class="panel-divider">', unsafe_allow_html=True)
    st.markdown('<div class="panel-label">Query</div>', unsafe_allow_html=True)

    no_doc = st.session_state.doc_id is None or not alive

    with st.form(key="ask_form", clear_on_submit=True):
        question = st.text_area(
            "Q", placeholder="Who is the consignee?\nWhat is the freight rate?",
            disabled=no_doc, label_visibility="collapsed", height=100,
        )
        ask_btn = st.form_submit_button(
            "Search Document →",
            disabled=no_doc, use_container_width=True, type="primary",
        )

    components.html("""
        <script>
        const doc = window.parent.document;
        const textareas = doc.querySelectorAll('textarea');
        textareas.forEach(ta => {
            if (ta.dataset.hasEnterListener) return;
            ta.dataset.hasEnterListener = "true";
            ta.addEventListener('keydown', function(e) {
                if (e.key === 'Enter' && !e.shiftKey) {
                    e.preventDefault();
                    const btn = doc.querySelector('button[kind="primaryFormSubmit"]');
                    if (btn) btn.click();
                }
            });
        });
        </script>
    """, height=0, width=0)

    if ask_btn and question.strip():
        with st.spinner("Retrieving…"):
            try:
                qv = embed_query(question.strip())
                resp = requests.post(
                    f"{get_url()}/ask",
                    json={
                        "doc_id":          st.session_state.doc_id,
                        "question":        question.strip(),
                        "query_embedding": qv.tolist(),
                    },
                    timeout=90,
                )
                if resp.status_code == 200:
                    result = resp.json()
                    st.session_state.chat_history.append({
                        "question":            question.strip(),
                        "answer":              result.get("answer", ""),
                        "confidence":          result.get("confidence", 0.0),
                        "sources":             result.get("sources", []),
                        "provider":            result.get("provider", "None"),
                        "logs":                result.get("logs", []),
                        "guardrail_triggered": result.get("guardrail_triggered", False),
                        "timestamp":           datetime.datetime.now().isoformat(),
                    })
                    st.session_state.active_tab = 0
                else:
                    st.error(f"Error: {resp.json().get('detail', resp.text)}")
            except requests.exceptions.ConnectionError:
                st.error(f"Cannot reach {get_url()}.")
            except Exception as e:
                st.error(f"Error: {e}")

    st.markdown('<hr class="panel-divider">', unsafe_allow_html=True)
    st.markdown('<div class="panel-label">Extract</div>', unsafe_allow_html=True)

    extract_btn = st.button(
        "Extract Shipment Data →",
        disabled=(st.session_state.doc_id is None or not alive),
        use_container_width=True,
    )

    if extract_btn:
        with st.spinner("Running extraction…"):
            try:
                resp = requests.post(
                    f"{get_url()}/extract",
                    json={"doc_id": st.session_state.doc_id},
                    timeout=120,
                )
                if resp.status_code == 200:
                    st.session_state.extract_result = resp.json()
                    st.session_state.active_tab     = 1
                else:
                    st.error(f"Error: {resp.json().get('detail', resp.text)}")
            except requests.exceptions.ConnectionError:
                st.error(f"Cannot reach {get_url()}.")
            except Exception as e:
                st.error(f"Error: {e}")

    if not st.session_state.doc_id:
        st.markdown("""
<div class="how-card">
  <div class="how-step"><span class="how-num">01</span>Upload a PDF, DOCX, or TXT logistics document</div>
  <div class="how-step"><span class="how-num">02</span>Ask questions — full chat history is preserved</div>
  <div class="how-step"><span class="how-num">03</span>Extract fields &amp; view the auto-generated shipment timeline</div>
</div>
""", unsafe_allow_html=True)

    if st.session_state.chat_history:
        st.markdown('<hr class="panel-divider">', unsafe_allow_html=True)
        if st.button("🗑  Clear Chat", use_container_width=True, type="secondary"):
            st.session_state.chat_history = []
            st.session_state.active_tab   = 0
            st.rerun()


# ════════════════════════════════════
#  RIGHT — Tabs
# ════════════════════════════════════
with right_col:

    n_msgs      = len(st.session_state.chat_history)
    has_extract = st.session_state.extract_result is not None

    tab_chat, tab_extract, tab_timeline = st.tabs([
        f"Chat  ({n_msgs})" if n_msgs else "Chat",
        "Extraction  ✓" if has_extract else "Extraction",
        "Timeline  ✓"   if has_extract else "Timeline",
    ])

    # ── Chat ─────────────────────────────────────────────────────────────────
    with tab_chat:
        if not st.session_state.chat_history:
            st.markdown("""
<div class="chat-empty">
  <div class="empty-glyph">Ask</div>
  <div class="empty-text">
    Upload a document and start asking questions.<br>
    Full conversation history is preserved<br>across extractions.
  </div>
</div>
""", unsafe_allow_html=True)
        else:
            st.markdown('<div class="chat-wrap">', unsafe_allow_html=True)
            for i, turn in enumerate(st.session_state.chat_history):
                bubble_cls = "chat-a-guard" if turn["guardrail_triggered"] else "chat-a"
                prefix     = "⚠ " if turn["guardrail_triggered"] else ""
                ts         = fmt_time(turn.get("timestamp", ""))
                meta = (
                    f'<div class="chat-meta">'
                    f'<span class="chat-ts">{ts}</span>'
                    f'{conf_html(turn["confidence"])}'
                    f'{provider_html(turn["provider"])}'
                    f'</div>'
                )
                st.markdown(f"""
<div>
  <div class="chat-q-row"><div class="chat-q">{turn["question"]}</div></div>
  <div class="chat-a-row">
    <div class="{bubble_cls}">{prefix}{turn["answer"]}{meta}</div>
  </div>
</div>
""", unsafe_allow_html=True)

                if turn["sources"] or turn["logs"]:
                    with st.expander(f"Sources ({len(turn['sources'])})  ·  #{i+1}", expanded=False):
                        for src in turn["sources"]:
                            pg = f"Page {src['page']}" if src.get("page") else "—"
                            st.markdown(
                                f'<div class="src-chunk">'
                                f'<div class="src-meta">chunk {src["chunk_index"]} · {pg} · {src["similarity"]:.0%}</div>'
                                f'{src["text"][:500]}{"…" if len(src["text"])>500 else ""}'
                                f'</div>',
                                unsafe_allow_html=True,
                            )
                        for log in turn["logs"]:
                            st.markdown(f'<div class="log-ln">{log}</div>', unsafe_allow_html=True)

                if i < len(st.session_state.chat_history) - 1:
                    st.markdown('<hr class="turn-divider">', unsafe_allow_html=True)

            st.markdown('</div>', unsafe_allow_html=True)

    # ── Extraction ────────────────────────────────────────────────────────────
    with tab_extract:
        if not has_extract:
            st.markdown("""
<div class="ext-empty">
  <div class="ext-empty-glyph">Extract</div>
  <div class="ext-empty-text">
    Click <b>Extract Shipment Data</b> in the left panel<br>
    to pull structured fields from the document.<br><br>
    Chat history is preserved while extraction runs.
  </div>
</div>
""", unsafe_allow_html=True)
        else:
            extract      = st.session_state.extract_result
            data         = extract.get("data", {})
            provider     = extract.get("provider", "None")
            ext_logs     = extract.get("logs", [])
            non_null     = sum(1 for v in data.values() if v is not None)
            total        = len(data)
            completeness = non_null / total if total else 0
            color = "#4ade80" if completeness >= 0.7 else "#f6a623" if completeness >= 0.45 else "#f87171"

            st.markdown(f"""
<div class="ext-stats">
  <div class="ext-stat">
    <div class="ext-stat-val" style="color:{color}">{completeness:.0%}</div>
    <div class="ext-stat-key">Completeness</div>
  </div>
  <div class="ext-stat">
    <div style="margin:4px 0 2px">{provider_html(provider)}</div>
    <div class="ext-stat-key">Provider</div>
  </div>
  <div class="ext-stat">
    <div class="ext-stat-val">{non_null}/{total}</div>
    <div class="ext-stat-key">Fields Found</div>
  </div>
</div>
""", unsafe_allow_html=True)
            st.progress(completeness)

            FIELD_LABELS = {
                "shipment_id": "shipment_id", "shipper": "shipper",
                "consignee": "consignee", "pickup_datetime": "pickup_datetime",
                "delivery_datetime": "delivery_datetime", "equipment_type": "equipment_type",
                "mode": "mode", "rate": "rate", "currency": "currency",
                "weight": "weight", "carrier_name": "carrier_name",
            }

            rows_html = ""
            for field, label in FIELD_LABELS.items():
                val = data.get(field)
                if val is not None:
                    if field == "rate" and isinstance(val, (int, float)):
                        rows_html += (
                            f'<div class="field-row">'
                            f'<span class="field-key">{label}</span>'
                            f'<span class="field-val-amber">${val:,.2f}</span>'
                            f'</div>'
                        )
                    else:
                        rows_html += (
                            f'<div class="field-row">'
                            f'<span class="field-key">{label}</span>'
                            f'<span class="field-val">{val}</span>'
                            f'</div>'
                        )
                else:
                    rows_html += (
                        f'<div class="field-row">'
                        f'<span class="field-key">{label}</span>'
                        f'<span class="field-null">null</span>'
                        f'</div>'
                    )

            st.markdown(f'<div class="field-card">{rows_html}</div>', unsafe_allow_html=True)

            with st.expander("Raw JSON", expanded=False):
                st.json(data)
            if ext_logs:
                with st.expander("Extraction Logs", expanded=False):
                    for log in ext_logs:
                        st.markdown(f'<div class="log-ln">{log}</div>', unsafe_allow_html=True)

    # ── Timeline (NEW) ────────────────────────────────────────────────────────
    with tab_timeline:
        if not has_extract:
            st.markdown("""
<div class="ext-empty">
  <div class="ext-empty-glyph">Timeline</div>
  <div class="ext-empty-text">
    Run <b>Extract Shipment Data</b> first.<br>
    The timeline will auto-build from extracted fields —<br>
    tracing the complete load journey from creation to delivery.
  </div>
</div>
""", unsafe_allow_html=True)
        else:
            data = st.session_state.extract_result.get("data", {})
            tl_html = render_timeline(data)

            st.markdown(f"""
<div class="timeline-card">
  <div class="timeline-header">
    <div class="timeline-title">Shipment Event Timeline</div>
    <div class="timeline-badge">Auto-generated</div>
  </div>
  {tl_html}
</div>
""", unsafe_allow_html=True)

            # Journey summary
            rate    = data.get("rate")
            cur     = data.get("currency", "")
            rate_s  = f"${rate:,.2f} {cur}".strip() if rate else "—"

            st.markdown(f"""
<div class="field-card">
  <div class="field-row">
    <span class="field-key">origin</span>
    <span class="field-val">{data.get('shipper') or '—'}</span>
  </div>
  <div class="field-row">
    <span class="field-key">destination</span>
    <span class="field-val">{data.get('consignee') or '—'}</span>
  </div>
  <div class="field-row">
    <span class="field-key">carrier</span>
    <span class="field-val">{data.get('carrier_name') or '—'}</span>
  </div>
  <div class="field-row">
    <span class="field-key">agreed rate</span>
    <span class="field-val-amber">{rate_s}</span>
  </div>
  <div class="field-row">
    <span class="field-key">equipment</span>
    <span class="field-val">{data.get('equipment_type') or '—'}</span>
  </div>
  <div class="field-row">
    <span class="field-key">weight</span>
    <span class="field-val">{data.get('weight') or '—'}</span>
  </div>
</div>
""", unsafe_allow_html=True)
